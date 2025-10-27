from fastapi import FastAPI, Request, UploadFile, Form
from fastapi.responses import HTMLResponse, RedirectResponse, FileResponse, PlainTextResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
import os
import shutil
import re
from docx import Document
import zipfile
import fitz  # PyMuPDF

app = FastAPI()

# Upload directory
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Static + templates
templates = Jinja2Templates(directory="templates")
app.mount("/static", StaticFiles(directory="static"), name="static")
app.mount("/uploads", StaticFiles(directory="uploads"), name="uploads")


@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    """Homepage showing all uploaded files"""
    files = sorted(os.listdir(UPLOAD_DIR))
    return templates.TemplateResponse("index.html", {"request": request, "files": files})


@app.post("/upload")
async def upload_file(file: UploadFile):
    """Upload a new file"""
    file_path = os.path.join(UPLOAD_DIR, file.filename)
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    return RedirectResponse(url="/", status_code=303)


@app.get("/view/{filename}", response_class=HTMLResponse)
def view_file(request: Request, filename: str):
    """View or read file content (supports txt, pdf, docx with images)"""
    file_path = os.path.join(UPLOAD_DIR, filename)
    if not os.path.exists(file_path):
        return RedirectResponse(url="/", status_code=303)

    ext = filename.split(".")[-1].lower()
    content = ""
    images = []

    try:
        if ext == "txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

        elif ext == "pdf":
            with fitz.open(file_path) as pdf_doc:
                extracted_text = []
                for page in pdf_doc:
                    text = page.get_text("text")
                    if text.strip():
                        extracted_text.append(text.strip())
                content = "\n\n".join(extracted_text)

        elif ext == "docx":
            doc = Document(file_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            content = "\n".join(paragraphs)

            img_dir = os.path.join(UPLOAD_DIR, f"{filename}_images")
            os.makedirs(img_dir, exist_ok=True)

            with zipfile.ZipFile(file_path, "r") as docx_zip:
                for file in docx_zip.namelist():
                    if file.startswith("word/media/") and (
                        file.endswith(".png") or file.endswith(".jpg") or file.endswith(".jpeg") or file.endswith(".gif")
                    ):
                        image_name = os.path.basename(file)
                        image_path = os.path.join(img_dir, image_name)
                        with open(image_path, "wb") as img_file:
                            img_file.write(docx_zip.read(file))
                        images.append(f"/uploads/{filename}_images/{image_name}")

            if not content.strip() and not images:
                content = "[This DOCX file appears empty or unsupported.]"

        else:
            content = "[Unsupported file format for preview.]"

        content = re.sub(r'\n\s*\n+', '\n\n', content.strip())

    except Exception as e:
        content = f"[Error reading file: {e}]"

    if not content.strip() and not images:
        content = "[No readable content found.]"

    return templates.TemplateResponse("view.html", {
        "request": request,
        "filename": filename,
        "content": content,
        "images": images
    })


@app.get("/download_text/{filename}", response_class=PlainTextResponse)
def download_text(filename: str):
    """Extract text and return as downloadable .txt file"""
    file_path = os.path.join(UPLOAD_DIR, filename)
    if not os.path.exists(file_path):
        return PlainTextResponse("File not found", status_code=404)

    ext = filename.split(".")[-1].lower()
    content = ""

    try:
        if ext == "txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

        elif ext == "pdf":
            with fitz.open(file_path) as pdf_doc:
                text_list = [page.get_text("text") for page in pdf_doc]
                content = "\n\n".join(text_list).strip()

        elif ext == "docx":
            doc = Document(file_path)
            content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

        else:
            return PlainTextResponse("Unsupported file type", status_code=400)

        if not content.strip():
            return PlainTextResponse("No readable text found.", status_code=204)

        # Create a downloadable response
        headers = {
            "Content-Disposition": f'attachment; filename="{os.path.splitext(filename)[0]}_extracted.txt"'
        }
        return PlainTextResponse(content, headers=headers)

    except Exception as e:
        return PlainTextResponse(f"Error extracting text: {e}", status_code=500)


@app.post("/update/{filename}")
def update_file(filename: str, new_content: str = Form(...)):
    """Update text content of a file"""
    file_path = os.path.join(UPLOAD_DIR, filename)
    ext = filename.split(".")[-1].lower()

    if ext not in ["txt"]:
        return RedirectResponse(url="/", status_code=303)

    with open(file_path, "w", encoding="utf-8") as f:
        f.write(new_content.strip())

    return RedirectResponse(url="/", status_code=303)


@app.get("/delete/{filename}")
def delete_file(filename: str):
    """Delete a file and its image folder"""
    file_path = os.path.join(UPLOAD_DIR, filename)
    img_dir = os.path.join(UPLOAD_DIR, f"{filename}_images")

    if os.path.exists(file_path):
        os.remove(file_path)

    if os.path.exists(img_dir):
        shutil.rmtree(img_dir)

    return RedirectResponse(url="/", status_code=303)
